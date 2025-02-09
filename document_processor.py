# document_processor.py (excerpt)
import os, re, json, csv
from typing import List, Dict, Any
import faiss
import numpy as np
import uuid
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS as LCFAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from PyPDF2 import PdfReader
import docx
from db_manager import insert_document, load_all_documents  # Import our DB functions

MODEL_NAME = "all-MiniLM-L6-v2"

class DocumentProcessor:
    def __init__(self, embedding_model_name: str = MODEL_NAME, chunk_size: int = 100, chunk_overlap: int = 20):
        self.embedding_model = HuggingFaceEmbeddings(model_name=embedding_model_name)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.documents: List[Document] = []
        embedding_dim = len(self.embedding_model.embed_query("test"))
        self.index = faiss.IndexFlatL2(embedding_dim)
        self.vector_store = LCFAISS(
            embedding_function=self.embedding_model,
            index=self.index,
            docstore=InMemoryDocstore(),
            index_to_docstore_id={}
        )

    def preprocess_text(self, text: str) -> str:
        text = text.strip()
        return re.sub(r'\s+', ' ', text)

    def chunk_text(self, text: str) -> List[str]:
        words = text.split()
        chunks = []
        start = 0
        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunks.append(" ".join(words[start:end]))
            start += self.chunk_size - self.chunk_overlap
        return chunks

    # Extraction methods for PDF, DOCX, CSV, etc. (same as before)
    def extract_text_from_pdf(self, filepath: str) -> str:
        text = ""
        try:
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF {filepath}: {e}")
        return text

    def extract_text_from_docx(self, filepath: str) -> str:
        try:
            doc = docx.Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error reading DOCX {filepath}: {e}")
            return ""

    def extract_text_from_csv(self, filepath: str) -> str:
        text = ""
        try:
            with open(filepath, "r", encoding="utf-8") as csvfile:
                reader = csv.reader(csvfile)
                for row in reader:
                    text += " | ".join(row) + "\n"
        except Exception as e:
            print(f"Error reading CSV {filepath}: {e}")
        return text

    def process_file(self, filepath: str, base_dir: str) -> None:
        ext = os.path.splitext(filepath)[1].lower()
        file_name = os.path.basename(filepath)
        metadata = {
            "title": file_name,
            "directory": base_dir,
            "source": base_dir,
            "tags": []
        }
        content = ""
        if ext == ".json":
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except Exception as e:
                print(f"Error reading JSON {filepath}: {e}")
                return
            if isinstance(data, list):
                for entry in data:
                    entry_metadata = metadata.copy()
                    entry_metadata.update({
                        "title": entry.get("title", file_name),
                        "region": entry.get("region", "Unknown"),
                        "source_url": entry.get("source_url", "Unknown"),
                        "tags": entry.get("tags", [])
                    })
                    description = self.preprocess_text(entry.get("description", ""))
                    self.add_document(description, entry_metadata, file_name)
            elif isinstance(data, dict):
                entry_metadata = metadata.copy()
                entry_metadata.update({
                    "title": data.get("title", file_name),
                    "region": data.get("region", "Unknown"),
                    "source_url": data.get("source_url", "Unknown"),
                    "tags": data.get("tags", [])
                })
                description = self.preprocess_text(data.get("description", ""))
                self.add_document(description, entry_metadata, file_name)
            return
        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        elif ext == ".pdf":
            content = self.extract_text_from_pdf(filepath)
        elif ext in [".doc", ".docx"]:
            content = self.extract_text_from_docx(filepath)
        elif ext == ".csv":
            content = self.extract_text_from_csv(filepath)
        else:
            print(f"Unsupported file type: {filepath}")
            return

        content = self.preprocess_text(content)
        self.add_document(content, metadata, file_name)

    def add_document(self, content: str, metadata: Dict[str, Any], file_name: str) -> None:
        chunks = self.chunk_text(content)
        for chunk in chunks:
            doc = Document(page_content=chunk, metadata=metadata)
            self.documents.append(doc)
            # Generate a unique ID for the document chunk.
            doc_id = str(uuid.uuid4())
            # Insert into the SQLite DB.
            from db_manager import insert_document  # Import here to avoid circular imports
            insert_document(doc_id, metadata.get("directory", "unknown"), file_name, chunk, metadata)
        # Note: Do not rebuild index on every addition.

    def build_index(self) -> None:
        if not self.documents:
            return
        self.vector_store.add_documents(documents=self.documents)
    
    def add_documents_from_directory(self, root_dir: str) -> None:
        file_count = 0
        for subdir, _, files in os.walk(root_dir):
            folder_name = os.path.basename(subdir)
            for file in files:
                filepath = os.path.join(subdir, file)
                self.process_file(filepath, base_dir=folder_name)
                file_count += 1
                print(f"Processed file: {filepath}")
        print(f"Total files processed: {file_count}")
        self.build_index()

    def load_documents_from_db(self) -> None:
        from db_manager import load_all_documents
        docs = load_all_documents()
        for doc_data in docs:
            metadata = doc_data["metadata"]
            doc = Document(page_content=doc_data["content"], metadata=metadata)
            self.documents.append(doc)
        self.build_index()

# For testing purposes:
if __name__ == "__main__":
    # Initialize DB first
    from db_manager import init_db
    init_db()
    dp = DocumentProcessor()
    # Alternatively, you could load from DB:
    dp.load_documents_from_db()
    results = dp.search("Tell me about superteam generally")
    print("Search results:")
    for doc in results:
        print(f"Title: {doc.metadata.get('title')}, Content: {doc.page_content[:100]}...")







